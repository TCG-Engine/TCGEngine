from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(__file__).with_name("TCGEngine_Remote_Frontend_Integration_Guide.docx")

NAVY = "1F4D78"
BLUE = "2E74B5"
INK = "172033"
MUTED = "5F6B7A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CODE_BG = "1F2937"
WHITE = "FFFFFF"
CONTENT_WIDTH_DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = tc_pr.first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                tc_pr.append(margins)
            for edge in ("top", "start", "bottom", "end"):
                elem = margins.find(qn(f"w:{edge}"))
                if elem is None:
                    elem = OxmlElement(f"w:{edge}")
                    margins.append(elem)
                elem.set(qn("w:w"), "80" if edge in ("top", "bottom") else "120")
                elem.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("TCGEngine remote frontend guide")
    set_font(run, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    add_page_field(footer)


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("TCGEngine")
    set_font(r, size=13, color=BLUE, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("Remote Frontend Integration Guide")
    set_font(r, size=27, color=INK, bold=True)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("A practical contract for connecting a custom client to an existing AzukiSim game hosted on petranaki.net")
    set_font(r, size=12, color=MUTED)

    callout = doc.add_table(rows=1, cols=1)
    set_table_geometry(callout, [CONTENT_WIDTH_DXA])
    cell = callout.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Integration model: ")
    set_font(r, size=10.5, color=NAVY, bold=True)
    r = p.add_run("the remote engine remains authoritative. Your frontend reads snapshots, renders them however it likes, and submits only legal intents back to the engine. It does not calculate game rules or mutate game state locally.")
    set_font(r, size=10.5, color=INK)

    add_heading(doc, "What this guide covers", 1)
    add_paragraph(doc, "This guide is for a friend building a replacement web frontend for an already-running AzukiSim match. It documents the production transport that the bundled browser client uses today: the long-poll state feed, the action endpoint, player and spectator authentication, decisions, chat, and the current delimiter-based snapshot format.")
    add_paragraph(doc, "It intentionally does not describe creating games, account management, deck builders, or internal admin/regression controls. Those are separate, sim- and deployment-specific surfaces.")


def add_heading(doc, text, level):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.left_indent = Inches(0.38)
    p.paragraph_format.first_line_indent = Inches(-0.19)
    r = p.add_run(text)
    set_font(r)
    return p


def add_code(doc, code):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, CODE_BG)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    for line_index, line in enumerate(code.strip("\n").splitlines()):
        if line_index:
            p.add_run("\n")
        r = p.add_run(line)
        set_font(r, name="Consolas", size=8.4, color=WHITE)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_font(r, size=9.5, color=NAVY, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def build_content(doc):
    add_heading(doc, "1. Start with a real seat or spectator link", 1)
    add_paragraph(doc, "The remote game is identified by its gameName, root folderPath, viewer identity, and authKey. Those values are normally present in the engine's NextTurn link. Treat the authKey as a bearer credential: anyone who has a player-seat link can issue actions for that seat.")
    add_table(doc, ["Viewer", "Request identity", "Visibility and authority"], [
        ("Player 1 or 2", "playerID=1 or playerID=2", "Receives that seat's private hand/temp-zone cards and may submit actions when the server permits it."),
        ("Spectator", "playerID=S; optional viewerPerspective=1 or 2", "Read-only. Private games require the spectator key; public-game spectator policy is game-specific."),
    ], [1800, 2200, 5360])
    add_paragraph(doc, "Use the exact deployment prefix supplied by the host. The examples below assume https://petranaki.net/TCGEngine, but the frontend should keep this as configuration because a host can mount TCGEngine under a different path.", "Use the exact deployment prefix")
    add_code(doc, "const config = {\n  baseUrl: 'https://petranaki.net/TCGEngine', // verify deployment path\n  folderPath: 'AzukiSim',\n  gameName: '37',\n  playerID: '1',       // or '2' / 'S'\n  viewerPerspective: '1', // meaningful for spectator views\n  authKey: '<seat-or-spectator-key>'\n};")

    add_heading(doc, "2. Transport overview", 1)
    add_table(doc, ["Purpose", "Endpoint", "Current wire format"], [
        ("Read live game state", "GET /AzukiSim/GetNextTurn.php", "Long-poll text response. Full board snapshots use <~> between fields and <|> between cards."),
        ("Submit a game action", "GET /ProcessInput.php", "Query-string action request. Optional responseFormat=json returns only an acknowledgement; pull state again after it."),
        ("Open a zone popup (optional)", "GET /GetPopupContent.php", "Legacy popup payload. A custom UI can usually render from the main snapshot instead."),
        ("Send chat (optional)", "GET /SubmitChat.php", "Plain-text acknowledgement; chat deltas also arrive in the main long-poll response."),
    ], [1840, 2680, 4840])
    add_paragraph(doc, "This is not a JSON REST board API. Do not attempt to write game files, infer legal moves, or apply client-side state transitions. Submit an intent, then accept the server's next snapshot as the source of truth.")

    add_heading(doc, "3. Long-polling the board", 1)
    add_paragraph(doc, "Call the game-specific GetNextTurn.php endpoint continuously. The server waits until the game or chat version changes, then returns a delta-bearing snapshot. The current AzukiSim loop checks roughly every 25 ms for up to 100 iterations (about 2.5 seconds) before returning KEEPALIVE.")
    add_code(doc, "const POLL = '/AzukiSim/GetNextTurn.php';\n\nasync function pollOnce(state) {\n  const u = new URL(config.baseUrl + POLL);\n  u.search = new URLSearchParams({\n    gameName: config.gameName,\n    playerID: config.playerID,\n    viewerPerspective: config.viewerPerspective,\n    authKey: config.authKey,\n    lastUpdate: String(state.lastUpdate),\n    lastChatVersion: String(state.lastChatVersion),\n    lastChatID: String(state.lastChatID),\n    windowWidth: String(window.innerWidth),\n    windowHeight: String(window.innerHeight)\n  });\n\n  const response = await fetch(u, { credentials: 'include' });\n  if (!response.ok) throw new Error(`poll failed: ${response.status}`);\n  return (await response.text()).trim();\n}")
    add_table(doc, ["Response", "Meaning", "Client action"], [
        ("KEEPALIVE or empty", "No board or chat update arrived during the hold period.", "Issue the next long-poll immediately; do not redraw."),
        ("CHATONLY<~>…", "Only chat changed.", "Parse the trailing chat JSON; keep the current board."),
        ("<number><~>…", "A newer game snapshot.", "Split into fields, reject stale update numbers, update chat, then render the full snapshot."),
        ("Invalid auth key / Invalid player", "Credentials or viewer identity failed validation.", "Stop polling and prompt for a fresh authorized link."),
    ], [1820, 2950, 4590])

    add_heading(doc, "4. Parsing the AzukiSim snapshot", 1)
    add_paragraph(doc, "For a normal board update, split the response on <~>. Field 0 is the monotonically increasing update number. Fields 1-14 are a raw player block. Fields 15-27 are global state. Fields 28-41 are the second raw player block. Optional trailing fields may contain bot-controller state, chat JSON, and frame-animation data; treat them as optional extensions rather than fixed board fields.")
    add_table(doc, ["Block offset", "Field", "How to use it"], [
        ("+1", "Deck", "Top card is CardBack with a count; do not reveal a deck order."),
        ("+2 to +4", "Discard, Hand, Temp zone", "Hand and temp-zone data are private to the matching seat; other viewers receive CardBack placeholders."),
        ("+5", "Global effects", "Reserved extension field; preserve it even if unused by your UI."),
        ("+6 to +8", "Garden, Alley, Gate", "Field zones. Each card carries its serialized object metadata."),
        ("+9 to +12", "Leader health, IKZ Area, IKZ pile, IKZ token", "Azuki-specific resource and state zones."),
        ("+13 to +14", "Decision queue, Versions", "Pending decisions drive prompts; versions are engine/replay history."),
    ], [1600, 2450, 5310])
    add_table(doc, ["Global index", "Value"], [
        ("15", "Turn number"), ("16", "First player"), ("17", "Turn player"), ("18", "Current phase"),
        ("19", "Phase parameters"), ("20", "Flash message"), ("21", "Decision queue variables JSON"),
        ("22", "Effect stack"), ("23", "Macro turn index"), ("24", "Unique-ID counter"),
        ("25", "Macro game index"), ("26", "Replay initial state"), ("27", "Replay commands"),
    ], [1800, 7560])
    add_paragraph(doc, "Perspective mapping: the wire response always emits the Player 1 block first and Player 2 block second. A player-facing UI should map the viewer's own raw block to my* and the other block to their*. A spectator can choose which player-facing orientation to render with viewerPerspective, but remains read-only.", "Perspective mapping:")

    add_heading(doc, "5. Card tuple and object metadata", 1)
    add_paragraph(doc, "A multi-card field is separated by <|>. Each card entry is a three-part legacy tuple: cardID, counter value, and serialized object JSON, separated by spaces. ClientRenderedCard replaces literal spaces in the JSON portion with underscores so that the tuple remains space-splittable. Parse the third token as JSON; preserve its string values exactly as supplied and do not use the counter slot as a general card-state model.")
    add_code(doc, "function parseCardTuple(entry) {\n  const [cardID, counterText = '0', objectJson = '-'] = entry.split(' ', 3);\n  return {\n    cardID,\n    counter: Number(counterText) || 0,\n    data: objectJson === '-' ? null : JSON.parse(objectJson)\n  };\n}\n\nfunction parseZone(field) {\n  if (!field) return [];\n  return field.split('<|>').filter(Boolean).map(parseCardTuple);\n}")
    add_paragraph(doc, "Important: the card object's mzID is the stable object handle used in subsequent actions. A card ID identifies a printed card; an mzID identifies that specific in-game object. Never replace an mzID with a card ID when submitting an action.")

    add_heading(doc, "6. Decisions and priority", 1)
    add_paragraph(doc, "The decision queue is the engine's prompt protocol. Do not show an arbitrary action picker while a decision is pending. Render the active decision first, submit exactly one response, and then wait for the next snapshot. The normal client uses the queue at the front of the acting player's decision queue and sends mode=DECISION, which ProcessInput maps to engine mode 100.")
    add_bullet(doc, "Use CurrentPhase, TurnPlayer, EffectStack, DecisionQueueVariables, and the viewer's decision queue to decide what to display; do not derive priority from local animations or timers.")
    add_bullet(doc, "If a decision supplies selectable cards, submit the chosen mzID in cardID. For a may-pass decision, submit cardID=PASS. Keep decisionIndex in the request because the existing client sends it, even though the active queue front remains authoritative.")
    add_bullet(doc, "For a multi-select decision, send the serialized result format expected by that decision type. Reuse the existing engine UI's serialization rules until a documented JSON decision endpoint is added.")
    add_code(doc, "async function answerDecision(decisionIndex, cardID) {\n  return submit({\n    mode: 'DECISION',\n    decisionIndex: String(decisionIndex),\n    cardID // e.g. 'myGarden-2' or 'PASS'\n  });\n}")

    add_heading(doc, "7. Submitting actions", 1)
    add_paragraph(doc, "All current game actions go through the shared ProcessInput.php endpoint. The request must carry gameName, playerID, authKey, folderPath, and mode. The server re-loads the game, verifies the seat authorization, rejects spectator actions, validates legality, executes the action, and advances the update cache. A successful acknowledgement is not a board snapshot; always poll again.")
    add_code(doc, "async function submit({ mode, cardID, decisionIndex, buttonInput, inputText }) {\n  const u = new URL(config.baseUrl + '/ProcessInput.php');\n  const q = new URLSearchParams({\n    gameName: config.gameName,\n    playerID: config.playerID,\n    authKey: config.authKey,\n    folderPath: config.folderPath,\n    mode: String(mode),\n    responseFormat: 'json'\n  });\n  if (cardID !== undefined) q.set('cardID', cardID);\n  if (decisionIndex !== undefined) q.set('decisionIndex', String(decisionIndex));\n  if (buttonInput !== undefined) q.set('buttonInput', buttonInput);\n  if (inputText !== undefined) q.set('inputText', inputText);\n  u.search = q;\n\n  const r = await fetch(u, { credentials: 'include' });\n  if (!r.ok) throw new Error(`action failed: ${r.status}`);\n  const ack = await r.json();\n  if (!ack.success) throw new Error(ack.message || 'action rejected');\n  return ack;\n}")
    add_table(doc, ["Intent", "Mode", "Typical cardID payload"], [
        ("Widget / activated ability", "10001", "<mzID>!<widgetType>!<action>"),
        ("Selectable card's zone action", "10002", "<mzID>!<action>!<comma-separated parameters>"),
        ("Decision response", "DECISION", "A chosen mzID, a serialized decision answer, or PASS"),
        ("Azuki normal pass", "10001", "myLeaderHealthSlot!CustomInput!Pass"),
    ], [2360, 1500, 5500])
    add_paragraph(doc, "Only submit payloads that the engine has exposed through the current snapshot, decision queue, or schema-driven UI metadata. The strings above document the existing client protocol; they are not permission to invent actions or bypass prereqs.")

    add_heading(doc, "8. Chat and optional popup support", 1)
    add_paragraph(doc, "Chat is already multiplexed into GetNextTurn. Include lastChatVersion and lastChatID in every poll. When the trailing chat payload advances, append only unseen message IDs. To send a message, call SubmitChat.php with gameName, playerID, authKey, folderPath, and chatText. Chat messages are sanitized and limited by the server.")
    add_code(doc, "const u = new URL(config.baseUrl + '/SubmitChat.php');\n+u.search = new URLSearchParams({\n  gameName: config.gameName, playerID: config.playerID,\n  authKey: config.authKey, folderPath: config.folderPath,\n  chatText: message\n});\nawait fetch(u, { credentials: 'include' });")
    add_paragraph(doc, "GetPopupContent.php is optional. It is the legacy mechanism for card-zone popups and returns a separate delimiter payload. A custom client should normally render discard, hand, and field zones directly from the main snapshot and only add popup support if it needs the existing macro-button behavior.")

    add_heading(doc, "9. Security and deployment rules", 1)
    add_bullet(doc, "Use HTTPS. Seat and spectator auth keys are credentials and currently travel in query strings because that is the existing engine contract.")
    add_bullet(doc, "Prefer serving the custom frontend from the same petranaki.net origin and deployment path. The engine endpoints do not constitute a documented cross-origin API; a separate origin needs an explicit reverse proxy or CORS policy controlled by the host.")
    add_bullet(doc, "Do not put full seat URLs in analytics, third-party error reporting, public screenshots, browser history exports, or referrer-bearing outbound links.")
    add_bullet(doc, "Never trust private-card data received under a different seat or perspective. Render the redacted CardBack values returned by the server.")
    add_bullet(doc, "Use one in-flight long-poll per game view. On network failure, retry with backoff; do not fan out parallel polling requests.")

    add_heading(doc, "10. Recommended client architecture", 1)
    add_number(doc, "Bootstrap from a validated game link and store the connection fields in memory, not local analytics or logs.")
    add_number(doc, "Start the long-poll loop with lastUpdate=0, lastChatVersion=0, and lastChatID=0.")
    add_number(doc, "Parse each full snapshot into an immutable view model. Map P1/P2 wire blocks into local/opponent perspective only after parsing.")
    add_number(doc, "Render cards, zones, counters, phase, priority, prompts, and chat from that view model. Do not mutate it optimistically after clicks.")
    add_number(doc, "Submit an action, wait for its acknowledgement, then let the next poll replace the view model. Surface server rejection messages as UI feedback.")
    add_number(doc, "Add a protocol adapter layer. Keep all delimiter parsing, mode strings, and action serialization in that one layer so a future JSON endpoint can replace it without rewriting the UI.")

    add_heading(doc, "11. Error handling checklist", 1)
    add_table(doc, ["Symptom", "Likely cause", "Correct behavior"], [
        ("401/403 or Invalid auth key", "Expired, wrong-seat, or private spectator key.", "Stop retries; require a fresh authorized link."),
        ("Snapshot update <= lastUpdate", "Delayed duplicate response.", "Ignore it and continue polling."),
        ("KEEPALIVE", "Expected long-poll timeout.", "Immediately issue the next poll."),
        ("Action acknowledgement says success=false", "Illegal action, wrong priority, or stale state.", "Show the message, then poll; never force local state."),
        ("Malformed delimiter payload", "Partial deploy, proxy error, or unexpected HTML/PHP output.", "Log a redacted sample, back off, and do not overwrite the current board."),
    ], [2200, 3400, 3760])

    add_heading(doc, "12. Source-of-truth files for future maintainers", 1)
    add_paragraph(doc, "This guide was verified against the current repository implementation. If the protocol changes, update the adapter and this document from these sources:")
    add_bullet(doc, "AzukiSim/GetNextTurn.php: board serialization, authentication, long-poll behavior, and Azuki-specific zone order.")
    add_bullet(doc, "AzukiSim/NextTurnRender.php: exact response-array indexes and player-perspective remapping used by the bundled frontend.")
    add_bullet(doc, "ProcessInput.php and Core/EngineActionRunner.php: action validation, DECISION normalization, and JSON acknowledgement behavior.")
    add_bullet(doc, "Core/jsInclude.js and Core/UILibraries20260717.js: current browser action serialization, decision UI, chat behavior, and zone rendering conventions.")
    add_bullet(doc, "Core/GameAuth.php and Core/ViewerIdentity.php: seat/spectator identity and authorization rules.")

    add_heading(doc, "Quick start reference", 1)
    add_code(doc, "1. GET  {base}/AzukiSim/GetNextTurn.php?...&lastUpdate=0\n2. If response is a snapshot: parse <~>, render it, save update/chat cursors.\n3. If the player chooses a legal server-exposed action: GET {base}/ProcessInput.php?...&responseFormat=json\n4. Do not update the board locally. Continue GET polling and render the next snapshot.\n5. On invalid auth, stop and obtain a new seat/spectator link.")


def main():
    doc = Document()
    configure_document(doc)
    add_title(doc)
    build_content(doc)
    doc.core_properties.title = "TCGEngine Remote Frontend Integration Guide"
    doc.core_properties.subject = "Custom frontend protocol for remotely hosted AzukiSim games"
    doc.core_properties.author = "TCGEngine"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
